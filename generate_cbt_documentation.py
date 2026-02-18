"""
CBT System Documentation Generator
Generates comprehensive Word documentation for the offline/online CBT system
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

def add_heading(doc, text, level):
    """Add a formatted heading"""
    heading = doc.add_heading(text, level=level)
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return heading

def add_paragraph(doc, text, bold=False, italic=False):
    """Add a formatted paragraph"""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.font.name = 'Calibri'
    if bold:
        run.bold = True
    if italic:
        run.italic = True
    return p

def add_bullet_list(doc, items):
    """Add a bulleted list"""
    for item in items:
        p = doc.add_paragraph(item, style='List Bullet')
        p.paragraph_format.left_indent = Inches(0.25)

def add_numbered_list(doc, items):
    """Add a numbered list"""
    for item in items:
        p = doc.add_paragraph(item, style='List Number')
        p.paragraph_format.left_indent = Inches(0.25)

def create_title_page(doc):
    """Create the title page"""
    title = doc.add_heading('Computer-Based Testing (CBT) System', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_paragraph()
    subtitle_run = subtitle.add_run('Offline & Online Examination Management')
    subtitle_run.font.size = Pt(16)
    subtitle_run.font.color.rgb = RGBColor(68, 114, 196)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    tagline = doc.add_paragraph()
    tagline_run = tagline.add_run('Complete Implementation Guide & Technical Overview')
    tagline_run.font.size = Pt(14)
    tagline_run.italic = True
    tagline.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_page_break()

def add_system_overview(doc):
    """Add system overview section"""
    add_heading(doc, '1. System Overview', 1)
    
    add_paragraph(doc, 
        'The CBT System is a comprehensive examination management platform designed to handle both '
        'online and offline examination scenarios. It provides seamless synchronization between '
        'connected and disconnected environments, ensuring examination continuity regardless of '
        'internet connectivity.'
    )
    
    doc.add_paragraph()
    add_heading(doc, 'Key Features', 2)
    add_bullet_list(doc, [
        'Dual-mode operation: Online and Offline examination delivery',
        'Automatic synchronization when connectivity is restored',
        'Role-based access control (Admin, Teacher, Student)',
        'Question bank management with subject and class-level organization',
        'Real-time exam monitoring and analytics',
        'Secure answer submission and validation',
        'Comprehensive marking and results management'
    ])
    
    doc.add_paragraph()
    add_heading(doc, 'System Architecture', 2)
    add_paragraph(doc,
        'The system follows a client-server architecture with Progressive Web App (PWA) capabilities:'
    )
    add_bullet_list(doc, [
        'Backend: Laravel 10.x with RESTful API (PHP)',
        'Frontend: React 18 with TypeScript and Tailwind CSS',
        'Database: MySQL with comprehensive relational schema',
        'Authentication: Laravel Sanctum token-based authentication',
        'Offline Storage: Browser LocalStorage and IndexedDB',
        'Service Workers: For offline caching and background sync'
    ])
    
    doc.add_page_break()

def add_online_exam_flow(doc):
    """Add online examination flow section"""
    add_heading(doc, '2. Online Examination Flow', 1)
    
    add_paragraph(doc,
        'When students have active internet connectivity, the system operates in online mode, '
        'providing real-time interaction with the server.'
    )
    
    doc.add_paragraph()
    add_heading(doc, 'Student Perspective', 2)
    add_numbered_list(doc, [
        'Login: Student authenticates using credentials',
        'Dashboard: View assigned exams with start/end times',
        'Exam Access: Click on available exam to begin',
        'Question Display: Questions load dynamically from server',
        'Answer Submission: Each answer saves immediately to server',
        'Real-time Sync: Progress bar and time tracker update live',
        'Completion: Submit all answers and receive confirmation',
        'Results: View scores when teacher releases marked exams'
    ])
    
    doc.add_paragraph()
    add_heading(doc, 'Teacher/Admin Perspective', 2)
    add_numbered_list(doc, [
        'Exam Creation: Set up exam with duration, subjects, classes',
        'Question Selection: Link questions from question bank',
        'Student Assignment: Allocate exam to specific students or classes',
        'Monitoring: Real-time dashboard shows active examinees',
        'Live Tracking: See which students started, in-progress, completed',
        'Marking: Review and grade open-ended questions',
        'Results Release: Publish results when ready',
        'Analytics: View performance statistics and reports'
    ])
    
    doc.add_paragraph()
    add_heading(doc, 'Technical Flow', 2)
    add_bullet_list(doc, [
        'API Endpoint: GET /api/exams/{id} retrieves exam details',
        'Questions: GET /api/exams/{id}/questions with randomization',
        'Answer Submit: POST /api/exam-attempts/{id}/answers',
        'Auto-save: Answers saved every 30 seconds or on answer change',
        'Validation: Server validates time limits and access permissions',
        'Completion: POST /api/exam-attempts/{id}/submit finalizes attempt'
    ])
    
    doc.add_page_break()

def add_offline_capture(doc):
    """Add offline capture mechanism section"""
    add_heading(doc, '3. Offline Examination Capture', 1)
    
    add_paragraph(doc,
        'The offline mode enables examinations to continue when internet connectivity is unavailable. '
        'All exam data and answers are stored locally on the device.'
    )
    
    doc.add_paragraph()
    add_heading(doc, 'Offline Preparation', 2)
    add_paragraph(doc, 'Before going offline, the system must pre-cache necessary data:')
    add_bullet_list(doc, [
        'Exam metadata (title, duration, instructions, subject, class)',
        'All exam questions with options',
        'Student authentication token',
        'Exam attempt record (created when student starts exam)',
        'Application assets (HTML, CSS, JavaScript files)',
        'Service worker registration for offline functionality'
    ])
    
    doc.add_paragraph()
    add_heading(doc, 'Offline Exam Taking', 2)
    add_numbered_list(doc, [
        'Detection: System detects loss of connectivity',
        'Mode Switch: UI indicates "Offline Mode" status',
        'LocalStorage: Answers stored in browser LocalStorage',
        'Question Navigation: Full navigation available from cached data',
        'Timer Tracking: Client-side timer continues counting',
        'Answer Recording: Each answer tagged with timestamp',
        'Submit Queue: Completion request queued for sync',
        'Validation: Client-side validation for required fields'
    ])
    
    doc.add_paragraph()
    add_heading(doc, 'Data Structure (Offline Storage)', 2)
    add_paragraph(doc, 'Answers are stored in JSON format in LocalStorage:')
    
    # Add code-like example
    code = doc.add_paragraph()
    code_run = code.add_run(
        '{\n'
        '  "attemptId": 123,\n'
        '  "examId": 45,\n'
        '  "studentId": 67,\n'
        '  "answers": [\n'
        '    {\n'
        '      "questionId": 1,\n'
        '      "answer": "A",\n'
        '      "timestamp": "2025-01-15T10:23:45Z"\n'
        '    },\n'
        '    {\n'
        '      "questionId": 2,\n'
        '      "answer": "Paris is the capital",\n'
        '      "timestamp": "2025-01-15T10:25:12Z"\n'
        '    }\n'
        '  ],\n'
        '  "submittedAt": "2025-01-15T10:45:00Z",\n'
        '  "syncStatus": "pending"\n'
        '}'
    )
    code_run.font.name = 'Courier New'
    code_run.font.size = Pt(9)
    code.paragraph_format.left_indent = Inches(0.5)
    
    doc.add_page_break()

def add_sync_process(doc):
    """Add synchronization process section"""
    add_heading(doc, '4. Synchronization Process', 1)
    
    add_paragraph(doc,
        'When connectivity is restored, the system automatically detects online status and initiates '
        'synchronization of all offline data to the server.'
    )
    
    doc.add_paragraph()
    add_heading(doc, 'Sync Trigger Mechanisms', 2)
    add_bullet_list(doc, [
        'Automatic Detection: Navigator.onLine event listener',
        'Periodic Checks: Heartbeat ping to server every 60 seconds',
        'Manual Retry: User-initiated sync button',
        'Service Worker: Background sync API for reliable delivery',
        'Resume on Page Load: Check for pending data on app startup'
    ])
    
    doc.add_paragraph()
    add_heading(doc, 'Sync Sequence', 2)
    add_numbered_list(doc, [
        'Connectivity Check: Verify server is reachable',
        'Authentication Refresh: Validate token, refresh if expired',
        'Data Retrieval: Load pending submissions from LocalStorage',
        'Validation: Check data integrity and completeness',
        'Transmission: POST answers to /api/exam-attempts/{id}/sync',
        'Server Processing: Backend validates and stores answers',
        'Confirmation: Server returns success/failure status',
        'Local Cleanup: Clear synced data from LocalStorage',
        'Error Handling: Retry failed submissions with exponential backoff',
        'User Notification: Display sync status to student'
    ])
    
    doc.add_paragraph()
    add_heading(doc, 'Sync API Endpoint', 2)
    add_paragraph(doc, 'Backend endpoint for offline sync:')
    
    api = doc.add_paragraph()
    api_run = api.add_run('POST /api/exam-attempts/{id}/sync')
    api_run.bold = True
    api_run.font.name = 'Courier New'
    
    add_paragraph(doc, 'Request Body:')
    request_body = doc.add_paragraph()
    rb_run = request_body.add_run(
        '{\n'
        '  "answers": [\n'
        '    {"question_id": 1, "answer": "A", "answered_at": "2025-01-15T10:23:45Z"},\n'
        '    {"question_id": 2, "answer": "Essay text", "answered_at": "2025-01-15T10:25:12Z"}\n'
        '  ],\n'
        '  "submitted_at": "2025-01-15T10:45:00Z",\n'
        '  "time_spent": 1320\n'
        '}'
    )
    rb_run.font.name = 'Courier New'
    rb_run.font.size = Pt(9)
    request_body.paragraph_format.left_indent = Inches(0.5)
    
    add_paragraph(doc, 'Response:')
    response = doc.add_paragraph()
    res_run = response.add_run(
        '{\n'
        '  "success": true,\n'
        '  "message": "Answers synced successfully",\n'
        '  "synced_count": 2,\n'
        '  "attempt_id": 123\n'
        '}'
    )
    res_run.font.name = 'Courier New'
    res_run.font.size = Pt(9)
    response.paragraph_format.left_indent = Inches(0.5)
    
    doc.add_page_break()

def add_conflict_resolution(doc):
    """Add conflict resolution section"""
    add_heading(doc, '5. Conflict Resolution & Edge Cases', 1)
    
    add_paragraph(doc,
        'The system handles various edge cases and conflicts that may arise during offline/online transitions.'
    )
    
    doc.add_paragraph()
    add_heading(doc, 'Time Limit Enforcement', 2)
    add_bullet_list(doc, [
        'Client Timer: Tracks exam duration on client side',
        'Server Validation: Backend checks submitted_at against exam end time',
        'Clock Skew: Server uses server time as authoritative source',
        'Grace Period: Optional 2-minute buffer for submission delays',
        'Late Submission Handling: Answers accepted but flagged as late',
        'Rejection: Submissions beyond grace period are rejected with error'
    ])
    
    doc.add_paragraph()
    add_heading(doc, 'Duplicate Submission Prevention', 2)
    add_bullet_list(doc, [
        'Attempt Status: Server tracks submission status (in-progress, submitted)',
        'Idempotency: Sync endpoint checks if attempt already submitted',
        'UUID Tracking: Each submission tagged with unique identifier',
        'Answer Merging: If new answers arrive, merge with existing (latest timestamp wins)',
        'Double Submit Block: Frontend disables submit button after first submission'
    ])
    
    doc.add_paragraph()
    add_heading(doc, 'Network Interruption During Sync', 2)
    add_bullet_list(doc, [
        'Partial Success: Track which answers successfully synced',
        'Retry Logic: Exponential backoff (retry after 10s, 30s, 60s, 300s)',
        'Persistence: Keep failed data in LocalStorage until confirmed',
        'Status Tracking: Mark each answer as "synced", "pending", "failed"',
        'Manual Intervention: Admin can manually sync from backend if needed'
    ])
    
    doc.add_paragraph()
    add_heading(doc, 'Student Switching Devices', 2)
    add_bullet_list(doc, [
        'Single Attempt Rule: One active attempt per student per exam',
        'Device Lock: Attempt tied to first device authentication token',
        'Resume Capability: Student can resume on same device',
        'Admin Override: Teacher can reset attempt if legitimate device change needed',
        'Security: Prevents exam sharing across multiple devices simultaneously'
    ])
    
    doc.add_paragraph()
    add_heading(doc, 'Browser Storage Limits', 2)
    add_bullet_list(doc, [
        'Storage Quota: Check available storage before caching exam',
        'Data Compression: Minimize JSON size for answers',
        'Progressive Sync: Sync answers in batches if large exam',
        'Cleanup: Remove old exam data after successful sync',
        'Fallback: Warn student to clear space if quota exceeded'
    ])
    
    doc.add_page_break()

def add_security_considerations(doc):
    """Add security section"""
    add_heading(doc, '6. Security & Integrity', 1)
    
    add_paragraph(doc,
        'The system implements multiple security layers to ensure exam integrity and prevent cheating.'
    )
    
    doc.add_paragraph()
    add_heading(doc, 'Authentication & Authorization', 2)
    add_bullet_list(doc, [
        'Token-Based Auth: Laravel Sanctum tokens with expiration',
        'Role Verification: Middleware checks user role (student/teacher/admin)',
        'Exam Access Control: Students only see assigned exams',
        'Time Window: Exams accessible only during scheduled time',
        'IP Restriction (Optional): Lock exam to specific network',
        'Single Session: Detect and prevent multiple concurrent sessions'
    ])
    
    doc.add_paragraph()
    add_heading(doc, 'Answer Tampering Prevention', 2)
    add_bullet_list(doc, [
        'Timestamp Validation: Server checks answer timestamps for chronological order',
        'Checksum: Optional hash of answer data to detect client-side modification',
        'Submit-Once: Mark attempt as submitted to prevent re-submission',
        'Question ID Verification: Ensure question IDs belong to the exam',
        'Answer Length Limits: Validate answer size to prevent storage abuse'
    ])
    
    doc.add_paragraph()
    add_heading(doc, 'Browser Tab Monitoring', 2)
    add_bullet_list(doc, [
        'Focus Tracking: Log when student switches tabs (optional)',
        'Full-Screen Mode: Encourage full-screen exam taking',
        'Warning System: Notify teacher if suspicious tab switching detected',
        'Auto-Submit: Option to force submit if student leaves tab multiple times'
    ])
    
    doc.add_paragraph()
    add_heading(doc, 'Data Privacy', 2)
    add_bullet_list(doc, [
        'HTTPS Only: All communication encrypted in transit',
        'Password Hashing: bcrypt for password storage',
        'Token Expiration: Sanctum tokens expire after inactivity',
        'LocalStorage Encryption: Optional encryption for offline data',
        'Audit Logs: Track all exam access and modifications',
        'GDPR Compliance: Student data deletion on request'
    ])
    
    doc.add_page_break()

def add_admin_configuration(doc):
    """Add admin configuration section"""
    add_heading(doc, '7. Administrator Configuration', 1)
    
    add_paragraph(doc,
        'Administrators can configure various aspects of the offline/online exam system.'
    )
    
    doc.add_paragraph()
    add_heading(doc, 'Exam Settings', 2)
    add_bullet_list(doc, [
        'Duration: Set exam time limit (minutes)',
        'Start/End Time: Schedule exam availability window',
        'Offline Mode: Enable/disable offline capability per exam',
        'Grace Period: Configure late submission tolerance (0-5 minutes)',
        'Auto-Submit: Force submit when timer expires',
        'Question Randomization: Shuffle question order per student',
        'Immediate Results: Show scores immediately or after teacher release'
    ])
    
    doc.add_paragraph()
    add_heading(doc, 'Question Bank Management', 2)
    add_bullet_list(doc, [
        'Subject Assignment: Link questions to subjects',
        'Class Level: Tag questions by educational level (e.g., JSS1, SSS2)',
        'Difficulty Levels: Mark questions as easy/medium/hard',
        'Mark Allocation: Set points per question',
        'Bulk Import: CSV upload for mass question creation',
        'Validation Rules: Enforce subject-class alignment'
    ])
    
    doc.add_paragraph()
    add_heading(doc, 'Student Allocation', 2)
    add_bullet_list(doc, [
        'Class-Based: Assign exam to entire class',
        'Individual: Select specific students manually',
        'Subject Filter: Only students taking the subject see exam',
        'Exam Access Management: Grant/revoke access as needed',
        'Bulk Allocation: CSV import for large-scale assignment'
    ])
    
    doc.add_paragraph()
    add_heading(doc, 'Monitoring Dashboard', 2)
    add_bullet_list(doc, [
        'Active Exams: View currently running examinations',
        'Student Status: See who started, in-progress, completed, not started',
        'Real-Time Alerts: Notifications for suspicious activity',
        'Sync Status: Track offline submissions awaiting sync',
        'Completion Rate: Percentage of assigned students who submitted',
        'Time Analytics: Average time spent per question'
    ])
    
    doc.add_paragraph()
    add_heading(doc, 'Marking & Results', 2)
    add_bullet_list(doc, [
        'Auto-Marking: Objective questions (MCQ, True/False) marked automatically',
        'Manual Marking: Review and grade subjective/essay questions',
        'Marking Summary: View pending vs completed marking tasks',
        'Grade Release: Publish results to students when ready',
        'Results Analytics: Performance statistics, pass/fail rates, grade distribution'
    ])
    
    doc.add_page_break()

def add_technical_implementation(doc):
    """Add technical implementation details"""
    add_heading(doc, '8. Technical Implementation Details', 1)
    
    add_paragraph(doc,
        'This section provides implementation specifics for developers and system administrators.'
    )
    
    doc.add_paragraph()
    add_heading(doc, 'Database Schema (Key Tables)', 2)
    
    # Exams table
    add_paragraph(doc, 'exams', bold=True)
    add_bullet_list(doc, [
        'id, title, description, subject_id, class_level',
        'duration (minutes), start_time, end_time',
        'total_marks, passing_marks',
        'randomize_questions (boolean)',
        'created_by (teacher/admin user_id)'
    ])
    
    doc.add_paragraph()
    
    # Bank questions table
    add_paragraph(doc, 'bank_questions', bold=True)
    add_bullet_list(doc, [
        'id, question_text, question_type (multiple_choice, essay, true_false)',
        'subject_id (required), class_level (required)',
        'options (JSON array), correct_answer',
        'marks, difficulty_level',
        'created_by, created_at, updated_at'
    ])
    
    doc.add_paragraph()
    
    # Exam questions pivot
    add_paragraph(doc, 'exam_questions (pivot)', bold=True)
    add_bullet_list(doc, [
        'id, exam_id, bank_question_id',
        'marks_override (optional, overrides bank question marks)',
        'order (for custom question sequencing)'
    ])
    
    doc.add_paragraph()
    
    # Exam attempts
    add_paragraph(doc, 'exam_attempts', bold=True)
    add_bullet_list(doc, [
        'id, exam_id, student_id',
        'started_at, submitted_at',
        'time_spent (seconds)',
        'status (in-progress, submitted, abandoned)',
        'score, total_marks',
        'offline_sync (boolean, true if synced from offline)'
    ])
    
    doc.add_paragraph()
    
    # Exam answers
    add_paragraph(doc, 'exam_answers', bold=True)
    add_bullet_list(doc, [
        'id, exam_attempt_id, question_id',
        'answer (text or option ID)',
        'is_correct (boolean, null for subjective)',
        'marks_awarded',
        'answered_at (timestamp)',
        'marked_by (teacher_id), marked_at'
    ])
    
    doc.add_paragraph()
    add_heading(doc, 'API Endpoints Summary', 2)
    
    endpoints = [
        ('POST /api/login', 'Authenticate user, return Sanctum token'),
        ('GET /api/exams', 'List exams (filtered by role and assignment)'),
        ('GET /api/exams/{id}', 'Get exam details with metadata'),
        ('GET /api/exams/{id}/questions', 'Get exam questions (randomized if enabled)'),
        ('POST /api/exam-attempts', 'Start exam attempt (record started_at)'),
        ('POST /api/exam-attempts/{id}/answers', 'Submit single answer'),
        ('POST /api/exam-attempts/{id}/submit', 'Finalize exam submission'),
        ('POST /api/exam-attempts/{id}/sync', 'Sync offline answers'),
        ('GET /api/results/analytics', 'Get results summary for admin/teacher'),
        ('GET /api/marking/exams', 'Get marking summary (pending/completed)')
    ]
    
    for endpoint, description in endpoints:
        p = doc.add_paragraph()
        endpoint_run = p.add_run(endpoint)
        endpoint_run.bold = True
        endpoint_run.font.name = 'Courier New'
        p.add_run(f' - {description}')
        p.paragraph_format.left_indent = Inches(0.25)
    
    doc.add_paragraph()
    add_heading(doc, 'Frontend Key Files', 2)
    add_bullet_list(doc, [
        'src/pages/student/TakeExam.tsx - Main exam taking interface',
        'src/hooks/useOfflineSync.ts - Custom hook for offline sync logic',
        'src/utils/offlineStorage.ts - LocalStorage management utilities',
        'src/services/examService.ts - API calls for exam operations',
        'public/service-worker.js - PWA service worker for offline caching'
    ])
    
    doc.add_paragraph()
    add_heading(doc, 'Backend Key Files', 2)
    add_bullet_list(doc, [
        'app/Http/Controllers/Api/ExamController.php - Exam CRUD operations',
        'app/Http/Controllers/Api/ExamAttemptController.php - Attempt tracking and sync',
        'app/Http/Controllers/Api/BankQuestionController.php - Question bank management',
        'app/Models/Exam.php - Exam model with relationships and calculations',
        'app/Models/ExamAttempt.php - Attempt model with scoring logic',
        'database/migrations/*_create_exams_table.php - Database schema definitions'
    ])
    
    doc.add_page_break()

def add_troubleshooting(doc):
    """Add troubleshooting section"""
    add_heading(doc, '9. Troubleshooting Common Issues', 1)
    
    doc.add_paragraph()
    add_heading(doc, 'Issue: Answers Not Syncing', 2)
    add_paragraph(doc, 'Symptoms: Student completed exam offline, but answers not appearing in backend.')
    add_paragraph(doc, 'Solutions:')
    add_bullet_list(doc, [
        'Check browser console for JavaScript errors',
        'Verify LocalStorage has answers: Open DevTools → Application → LocalStorage',
        'Confirm internet connectivity: Look for "Online" status indicator',
        'Check token expiration: Student may need to re-login',
        'Manually trigger sync: Use admin sync utility if available',
        'Review server logs: Look for sync endpoint errors (401, 422, 500)',
        'Verify exam attempt status: Ensure attempt not already marked as submitted'
    ])
    
    doc.add_paragraph()
    add_heading(doc, 'Issue: Time Limit Exceeded Error', 2)
    add_paragraph(doc, 'Symptoms: Student submits on time but receives "time expired" error.')
    add_paragraph(doc, 'Solutions:')
    add_bullet_list(doc, [
        'Check server time vs client time: Clock skew can cause issues',
        'Verify grace period setting: Adjust if network lag is common',
        'Review submitted_at timestamp: Ensure client sends correct time',
        'Check timezone configuration: Backend and frontend must align',
        'Admin override: Manually accept late submission if legitimate'
    ])
    
    doc.add_paragraph()
    add_heading(doc, 'Issue: Questions Not Loading Offline', 2)
    add_paragraph(doc, 'Symptoms: Student loses connection and sees blank exam page.')
    add_paragraph(doc, 'Solutions:')
    add_bullet_list(doc, [
        'Ensure exam was accessed online first (cache must be populated)',
        'Check service worker registration: DevTools → Application → Service Workers',
        'Verify cache storage: Look for exam data in Cache or IndexedDB',
        'Clear and re-cache: Have student reload exam while online',
        'Browser compatibility: Some browsers have limited offline support',
        'Storage quota: Confirm browser has sufficient storage space'
    ])
    
    doc.add_paragraph()
    add_heading(doc, 'Issue: Duplicate Submissions', 2)
    add_paragraph(doc, 'Symptoms: Same exam attempt submitted multiple times.')
    add_paragraph(doc, 'Solutions:')
    add_bullet_list(doc, [
        'Implement idempotency key: Tag each submission with unique UUID',
        'Check submission status: Backend should reject if already submitted',
        'Disable submit button: Frontend should prevent multiple clicks',
        'Clear LocalStorage after sync: Remove data once confirmed synced',
        'Admin cleanup: Delete duplicate attempts, keep latest one'
    ])
    
    doc.add_paragraph()
    add_heading(doc, 'Issue: Poor Offline Performance', 2)
    add_paragraph(doc, 'Symptoms: Exam interface slow or unresponsive when offline.')
    add_paragraph(doc, 'Solutions:')
    add_bullet_list(doc, [
        'Optimize data structure: Reduce JSON size of cached questions',
        'Limit question images: Use lower resolution or defer loading',
        'Minimize LocalStorage writes: Batch answer saves',
        'Use IndexedDB for large data: Better performance than LocalStorage',
        'Reduce UI complexity: Simplify rendering during offline mode',
        'Browser profiling: Use DevTools Performance tab to identify bottlenecks'
    ])
    
    doc.add_page_break()

def add_best_practices(doc):
    """Add best practices section"""
    add_heading(doc, '10. Best Practices & Recommendations', 1)
    
    doc.add_paragraph()
    add_heading(doc, 'For Administrators', 2)
    add_bullet_list(doc, [
        'Test Offline Mode: Conduct trial exam with teachers before deploying to students',
        'Monitor Sync Status: Regularly check for failed syncs and resolve promptly',
        'Set Realistic Time Limits: Factor in potential network delays',
        'Enable Grace Period: 2-3 minute buffer for submission delays',
        'Regular Backups: Export exam attempts daily during examination period',
        'Clear Old Data: Archive completed exams to maintain system performance',
        'User Training: Educate students on offline mode before high-stakes exams',
        'Stable Environment: Ensure server uptime during scheduled exam windows'
    ])
    
    doc.add_paragraph()
    add_heading(doc, 'For Teachers', 2)
    add_bullet_list(doc, [
        'Advance Preparation: Create exams at least 24 hours before start time',
        'Question Review: Verify all questions have correct answers and marks',
        'Student Allocation: Confirm all students assigned before exam starts',
        'Subject-Class Alignment: Ensure bank questions match exam subject/class',
        'Mark Allocation: Verify total marks add up correctly',
        'Timely Marking: Complete subjective marking within 48 hours',
        'Communication: Notify students of results release',
        'Feedback: Review exam analytics to improve future assessments'
    ])
    
    doc.add_paragraph()
    add_heading(doc, 'For Students', 2)
    add_bullet_list(doc, [
        'Pre-load Exam: Access exam page while online before connection loss',
        'Stable Browser: Use Chrome, Firefox, or Edge (avoid older browsers)',
        'Single Tab: Keep exam in one tab, avoid switching',
        'Manual Save: Click save button periodically even though auto-save is enabled',
        'Monitor Timer: Track remaining time, submit before expiration',
        'Connectivity Check: Note if "Offline Mode" appears on screen',
        'Stay Calm: Offline answers will sync automatically when online',
        'Technical Issues: Contact teacher immediately if problems arise'
    ])
    
    doc.add_paragraph()
    add_heading(doc, 'For Developers', 2)
    add_bullet_list(doc, [
        'Error Logging: Implement comprehensive logging for sync failures',
        'Retry Strategy: Exponential backoff with maximum retry limit',
        'Data Validation: Validate all inputs on both client and server',
        'Idempotency: Ensure sync endpoints can handle duplicate requests',
        'Performance Testing: Load test with realistic concurrent users',
        'Security Audits: Regularly review authentication and authorization',
        'Code Documentation: Maintain clear comments for offline logic',
        'Version Control: Tag releases and maintain changelog',
        'Monitoring: Set up alerts for sync failures and errors',
        'User Feedback: Collect and address usability issues promptly'
    ])
    
    doc.add_page_break()

def add_conclusion(doc):
    """Add conclusion section"""
    add_heading(doc, 'Conclusion', 1)
    
    add_paragraph(doc,
        'The CBT System provides a robust and reliable platform for conducting examinations in both '
        'online and offline scenarios. The synchronization mechanism ensures that student answers are '
        'preserved and submitted even when connectivity is intermittent or unavailable.'
    )
    
    doc.add_paragraph()
    add_paragraph(doc,
        'Key strengths of the system include:'
    )
    add_bullet_list(doc, [
        'Seamless transition between online and offline modes',
        'Comprehensive security and integrity measures',
        'Role-based access control for admin, teacher, and student',
        'Flexible exam configuration and question bank management',
        'Real-time monitoring and detailed analytics',
        'Automatic and manual synchronization capabilities',
        'Conflict resolution for edge cases',
        'User-friendly interface for all stakeholders'
    ])
    
    doc.add_paragraph()
    add_paragraph(doc,
        'By following the best practices outlined in this document, institutions can successfully '
        'deploy and operate the CBT system to facilitate secure, efficient, and fair examinations '
        'regardless of internet connectivity constraints.'
    )
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    footer = doc.add_paragraph()
    footer_run = footer.add_run('For technical support and updates, please contact your system administrator.')
    footer_run.italic = True
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER

def main():
    """Main function to generate the document"""
    print("Starting CBT Documentation Generation...")
    
    # Create document
    doc = Document()
    
    # Set default style
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    
    # Create title page
    print("Creating title page...")
    create_title_page(doc)
    
    # Add sections
    print("Adding system overview...")
    add_system_overview(doc)
    
    print("Adding online exam flow...")
    add_online_exam_flow(doc)
    
    print("Adding offline capture mechanism...")
    add_offline_capture(doc)
    
    print("Adding synchronization process...")
    add_sync_process(doc)
    
    print("Adding conflict resolution...")
    add_conflict_resolution(doc)
    
    print("Adding security considerations...")
    add_security_considerations(doc)
    
    print("Adding admin configuration...")
    add_admin_configuration(doc)
    
    print("Adding technical implementation...")
    add_technical_implementation(doc)
    
    print("Adding troubleshooting...")
    add_troubleshooting(doc)
    
    print("Adding best practices...")
    add_best_practices(doc)
    
    print("Adding conclusion...")
    add_conclusion(doc)
    
    # Save document
    output_path = 'docs/CBT_Offline_Online_Complete_Guide.docx'
    doc.save(output_path)
    print(f"\n✓ Documentation generated successfully: {output_path}")
    print(f"Total pages: {len(doc.sections)} section(s)")
    print(f"Total paragraphs: {len(doc.paragraphs)}")

if __name__ == '__main__':
    main()
